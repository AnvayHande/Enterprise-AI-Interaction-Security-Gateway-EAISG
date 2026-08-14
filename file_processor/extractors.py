import io
import pdfplumber
import docx
import openpyxl
from fastapi import HTTPException
import logging

logger = logging.getLogger(__name__)

class FileExtractor:
    @staticmethod
    def extract_text(file_content: bytes, mime_type: str) -> str:
        try:
            if mime_type == "text/plain" or mime_type == "application/json":
                # Assuming utf-8 for MVP, with fallback replacing errors
                return file_content.decode('utf-8', errors='replace')
                
            elif mime_type == "application/pdf":
                text_parts = []
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                return "\n".join(text_parts)
                
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = docx.Document(io.BytesIO(file_content))
                text_parts = [para.text for para in doc.paragraphs]
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text:
                                text_parts.append(cell.text)
                return "\n".join(text_parts)
                
            elif mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                wb = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
                text_parts = []
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows(values_only=True):
                        row_texts = [str(cell) for cell in row if cell is not None]
                        if row_texts:
                            text_parts.append(" ".join(row_texts))
                return "\n".join(text_parts)
                
            else:
                raise HTTPException(status_code=415, detail=f"No extractor implemented for {mime_type}")
                
        except Exception as e:
            logger.error(f"Failed to extract content: {str(e)}")
            raise HTTPException(status_code=500, detail="Failed to safely extract file content.")
